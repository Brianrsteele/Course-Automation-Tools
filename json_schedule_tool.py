import json
import datetime
from datetime import timedelta
from docx import Document
from docx.shared import Inches


# Creating a semester schedule by scraping the rctc web page is proving to be too fragile.
# The page changes every year and the scraping methods need to be re-written.
# It is making more sense to format the dates into a JSON document by hand.

# Brian Steele - 7/21/22

# json formated file to read dates from
# must be updated before running
DATES_FILE = "2024-25.json"
YEAR = "2025"
SEMESTER = "Spring"
COURSE = "ART 2280"
SECTION = "31"

# these need to be the full day name and capitalized
DAYS_OF_WEEK = ["Sunday"]
FILE_NAME = COURSE + "-" + SECTION + " " + SEMESTER + " " + YEAR


def main():
    # read the json dates file the years holidays
    # and important dates recorded
    f = open(DATES_FILE)
    data = json.load(f)

    # find the data for important dates in the
    # data dictionary
    cal = data["important_dates"]

    print(cal)

    # add [0] at the end to get just the dictionary, not the list which contains
    # a dictionary
    semester = [
        sem for sem in cal if sem["year"] == YEAR and sem["semester"] == SEMESTER
    ][0]

    # get rid of year and semester fields
    # and keys for values
    important_dates = {
        item: value
        for (item, value) in semester.items()
        if item != "year" and item != "semester"
    }

    start_date = ""
    end_date = ""

    # find begin and end dates in the list of important dates
    for date in important_dates:
        if "begin" in important_dates[date].lower():
            start_date = date
        if "end" in important_dates[date].lower():
            end_date = date

    # deal with not find start and end dates
    if not start_date:
        raise Exception("Cound not find a start date")
    if not end_date:
        raise Exception("Cound not find an end date")

    # create datetime.datetime object for start and end dates
    start_date = return_datetime_object(start_date)
    end_date = return_datetime_object(end_date)

    # create a list of date objects for all the meeting dates
    meeting_days = return_schedule(start_date, end_date)

    schedule = []

    # create a list of dates with holiday information
    for date in meeting_days:
        # need to know the date string to compare with the important_dates...
        # this will likely be a problem if the json data isn't in the format
        # 'August 22, 2022'.

        date_string = date.strftime("%B %d, %G")
        date_string_with_weekday = date.strftime("%A, %B %d, %G")

        if date.strftime("%A") in DAYS_OF_WEEK:
            if date_string in important_dates:
                schedule.append(
                    (date_string_with_weekday, important_dates[date_string])
                )
            else:
                schedule.append((date_string_with_weekday, ""))

    export_md(schedule)
    export_docx(schedule)


def return_datetime_object(the_date):
    """
    Turns a date in the format "August 22, 2022" into
    a datetime.datetime object.

    @param: string with a date in 'August 22, 2022' format
    @returns: a datetime.datetime object
    """
    months = {
        "jan": 1,
        "feb": 2,
        "mar": 3,
        "apr": 4,
        "may": 5,
        "jun": 6,
        "jul": 7,
        "aug": 8,
        "sep": 9,
        "oct": 10,
        "nov": 11,
        "dec": 12,
    }
    year = int(the_date.split(",")[1].strip())
    month = the_date.split()[0].strip().lower()[:3]
    month = months[month]
    day = int(the_date.split()[1].split(",")[0].strip())
    return datetime.datetime(year, month, day)


def return_schedule(start, end):
    """
    Creates a list of meeting dates for a semester.

    @param: start - a datetime.datetime object for the start of the semester
    @returns: list of datetime.datetime objects with the semester class dates
    """
    one_date = datetime.timedelta(days=1)
    schedule = [
        start + datetime.timedelta(days=x) for x in range((end - start).days + 1)
    ]
    return schedule


def export_md(schedule):
    """
    Exports a markdown version of the schedule.

    @param: schedule - a list object with dates.
    @returns: None - saves a markdown file with the dates in markdown form.
    """
    document = (
        "# Important Dates for "
        + COURSE
        + "-"
        + SECTION
        + ", "
        + SEMESTER
        + " "
        + YEAR
        + "#\n"
    )
    for date in schedule:
        document += "## " + date[0] + " ##\n\n"
        document += "* " + date[1] + "\n\n"

    md_file = FILE_NAME + ".md"
    with open(md_file, "x") as file:
        file.write(document)


def export_docx(schedule):
    """
    Saves a word document with an outline of the course dates

    @param: schedule - a list with the dates for the semester.
    """
    important_dates = Document()
    title = important_dates.add_heading(
        "Important Dates for " + COURSE + "-" + SECTION + ", " + SEMESTER + " " + YEAR,
        0,
    )
    for date in schedule:
        important_dates.add_heading(date[0])
        important_dates.add_paragraph(date[1], style="List Bullet")

    important_dates.save(FILE_NAME + ".docx")


if __name__ == "__main__":
    main()
